"""可见光到红外图像生成页面：单张/批量完整工作流。"""

import os
import subprocess
import sys
import time
from pathlib import Path

from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont
from PyQt5.QtWidgets import (
    QCheckBox,
    QComboBox,
    QFileDialog,
    QFrame,
    QGridLayout,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QProgressBar,
    QPushButton,
    QRadioButton,
    QScrollArea,
    QSizePolicy,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from core.ir_generation import generate_ir_algorithm_1, generate_ir_algorithm_2
from core.metrics import calculate_metrics, get_lipips_error
from core.utils import (
    IMAGE_EXTENSIONS,
    create_comparison_image,
    find_matching_label,
    load_image,
    save_image_to_path,
    scan_image_files,
    show_image_on_label,
)


DEFAULT_OUTPUT_DIR = "D:/Generated_IR"
IMAGE_FILTER = "图像文件 (*.jpg *.jpeg *.png *.bmp *.tif *.tiff);;所有文件 (*.*)"
def _algorithm_key(algorithm_id):
    return "DFSMamba" if algorithm_id == 1 else "PhysMamba"


def _safe_prefix(prefix):
    prefix = (prefix or "").strip()
    return "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in prefix)


def _safe_dir_name(name, fallback="unnamed"):
    name = (name or "").strip()
    cleaned = "".join(ch if ch.isalnum() or ch in ("-", "_", ".") else "_" for ch in name)
    cleaned = cleaned.strip(" ._")
    return cleaned or fallback


def _unique_dir(base_dir, preferred_name):
    """返回固定结果目录；同名目录复用，内部结果文件覆盖。"""
    return str(Path(base_dir) / _safe_dir_name(preferred_name))


def _result_stem(input_path, algorithm_key, prefix=""):
    stem = Path(input_path).stem
    prefix = _safe_prefix(prefix)
    if prefix:
        stem = f"{prefix}_{stem}"
    return f"{stem}_{algorithm_key}"


def _metric_text(value):
    if value is None or value == "":
        return ""
    return f"{float(value):.4f}"


def _call_generation_algorithm(algorithm_id, input_image, device_preference):
    if algorithm_id == 1:
        return generate_ir_algorithm_1(input_image, device_preference=device_preference)
    return generate_ir_algorithm_2(input_image, device_preference=device_preference)


def _write_params_file(params, params_path):
    os.makedirs(os.path.dirname(params_path), exist_ok=True)
    with open(params_path, "w", encoding="utf-8") as handle:
        for key, value in params.items():
            handle.write(f"{key}: {value}\n")


def _average_metric(records, key):
    values = []
    for record in records:
        raw = record.get(key, "")
        if raw in ("", None):
            continue
        try:
            values.append(float(raw))
        except (TypeError, ValueError):
            continue
    if not values:
        return ""
    return f"{sum(values) / len(values):.4f}"


def _record_report_line(key, value):
    return f"{key}: {value}"


def _write_summary_file(summary, records, summary_path):
    os.makedirs(os.path.dirname(summary_path), exist_ok=True)
    lines = [
        "红外图像生成任务总结",
        "",
        f"输入文件夹: {summary.get('input_dir', '')}",
        f"输出文件夹: {summary.get('output_dir', '')}",
        f"标签文件夹: {summary.get('label_dir', '')}",
        f"生成算法: {summary.get('algorithm', '')}",
        f"运行设备: {summary.get('device', '')}",
        f"开始时间: {summary.get('start_time', '')}",
        f"结束时间: {summary.get('end_time', '')}",
        f"任务状态: {summary.get('task_status', '')}",
        f"总图片数: {summary.get('total', 0)}",
        f"已处理数量: {summary.get('completed', 0)}",
        f"成功数量: {summary.get('success', 0)}",
        f"失败数量: {summary.get('failed', 0)}",
        f"总耗时: {summary.get('total_elapsed', 0):.3f} s",
        f"平均单张耗时: {summary.get('avg_elapsed', 0):.3f} s",
        f"平均 PSNR: {_average_metric(records, 'psnr')}",
        f"平均 SSIM: {_average_metric(records, 'ssim')}",
        f"平均 L1: {_average_metric(records, 'l1')}",
        f"平均 LIPIPS: {_average_metric(records, 'lipips')}",
    ]
    with open(summary_path, "w", encoding="utf-8") as handle:
        handle.write("\n".join(lines))


def _write_batch_results_txt(summary, records, params, txt_path):
    """把一次批处理的参数、汇总和逐图评估结果写入同一个 TXT 文件。"""
    os.makedirs(os.path.dirname(txt_path), exist_ok=True)
    lines = [
        "批量红外图像生成结果",
        "",
        "一、运行参数",
    ]
    for key, value in params.items():
        lines.append(f"{key}: {value}")

    lines.extend(
        [
            "",
            "二、汇总统计",
            f"输入文件夹: {summary.get('input_dir', '')}",
            f"输出文件夹: {summary.get('output_dir', '')}",
            f"标签文件夹: {summary.get('label_dir', '')}",
            f"生成算法: {summary.get('algorithm', '')}",
            f"运行设备: {summary.get('device', '')}",
            f"开始时间: {summary.get('start_time', '')}",
            f"结束时间: {summary.get('end_time', '')}",
            f"任务状态: {summary.get('task_status', '')}",
            f"总图片数: {summary.get('total', 0)}",
            f"已处理数量: {summary.get('completed', 0)}",
            f"成功数量: {summary.get('success', 0)}",
            f"失败数量: {summary.get('failed', 0)}",
            f"总耗时: {summary.get('total_elapsed', 0):.3f} s",
            f"平均单张耗时: {summary.get('avg_elapsed', 0):.3f} s",
            f"平均 PSNR: {_average_metric(records, 'psnr')}",
            f"平均 SSIM: {_average_metric(records, 'ssim')}",
            f"平均 L1: {_average_metric(records, 'l1')}",
            f"平均 LIPIPS: {_average_metric(records, 'lipips')}",
            "",
            "三、逐张结果",
        ]
    )

    for index, record in enumerate(records, 1):
        lines.extend(
            [
                "",
                f"[{index}] {record.get('image_name', '')}",
                f"input_path: {record.get('input_path', '')}",
                f"label_path: {record.get('label_path', '')}",
                f"output_path: {record.get('output_path', '')}",
                f"compare_path: {record.get('compare_path', '')}",
                f"algorithm: {record.get('algorithm', '')}",
                f"device: {record.get('device', '')}",
                f"status: {record.get('status', '')}",
                f"error_message: {record.get('error_message', '')}",
                f"inference_time: {record.get('inference_time', '')}",
                f"psnr: {record.get('psnr', '')}",
                f"ssim: {record.get('ssim', '')}",
                f"l1: {record.get('l1', '')}",
                f"lipips: {record.get('lipips', '')}",
            ]
        )

    with open(txt_path, "w", encoding="utf-8") as handle:
        handle.write("\n".join(lines))


def _write_single_txt_report(record, params, txt_path):
    os.makedirs(os.path.dirname(txt_path), exist_ok=True)
    lines = [
        "单张红外图像生成报告",
        "",
        f"输入图像: {record.get('input_path', '')}",
        f"标签图像: {record.get('label_path', '')}",
        f"生成图像: {record.get('output_path', '')}",
        f"对比图像: {record.get('compare_path', '')}",
        f"生成算法: {record.get('algorithm', '')}",
        f"运行设备: {record.get('device', '')}",
        f"状态: {record.get('status', '')}",
        f"错误信息: {record.get('error_message', '')}",
        f"推理耗时: {record.get('inference_time', '')} s",
        f"PSNR: {record.get('psnr', '')}",
        f"SSIM: {record.get('ssim', '')}",
        f"L1: {record.get('l1', '')}",
        f"LIPIPS: {record.get('lipips', '')}",
        "",
        "运行参数",
    ]
    for key, value in params.items():
        lines.append(f"{key}: {value}")
    with open(txt_path, "w", encoding="utf-8") as handle:
        handle.write("\n".join(lines))


def _try_write_docx_report(title, lines, docx_path):
    from docx import Document

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    document = Document()
    document.add_heading(title, 0)
    for line in lines:
        document.add_paragraph(str(line))
    document.save(docx_path)


class GenerationWorker(QThread):
    """单张红外图像生成工作线程。"""

    log_message = pyqtSignal(str)
    progress_changed = pyqtSignal(object)
    preview_ready = pyqtSignal(object, object, object, object)
    result_ready = pyqtSignal(object)
    error_occurred = pyqtSignal(str)

    def __init__(self, input_path, label_path, options):
        super().__init__()
        self.input_path = input_path
        self.label_path = label_path
        self.options = options
        self.cancel_requested = False

    def request_cancel(self):
        self.cancel_requested = True

    def run(self):
        start_readable = time.strftime("%Y-%m-%d %H:%M:%S")
        start_clock = time.time()
        algorithm_id = self.options["algorithm_id"]
        algorithm_key = self.options["algorithm_key"]
        output_root = self.options["output_dir"]
        single_root = os.path.join(output_root, "单张文件")
        image_folder_name = Path(self.input_path).stem
        run_dir = _unique_dir(single_root, image_folder_name)
        generated_dir = os.path.join(run_dir, "generated")
        compare_dir = os.path.join(run_dir, "compare")
        reports_dir = os.path.join(run_dir, "reports")

        try:
            for folder in (generated_dir, compare_dir, reports_dir):
                os.makedirs(folder, exist_ok=True)
        except Exception as exc:
            self.error_occurred.emit(f"输出目录无法创建: {exc}")
            return

        record = self._empty_record()
        params = self._params(run_dir)

        try:
            self.log_message.emit("[RUN] 启动单张生成任务")
            self.log_message.emit(f"[INFO] 当前处理模式: 单张")
            self.log_message.emit(f"[INFO] 输入路径: {self.input_path}")
            self.log_message.emit(f"[INFO] 输出路径: {run_dir}")
            self.log_message.emit(f"[INFO] 使用算法: {algorithm_key} | 使用设备: {self.options['device']}")
            self._emit_progress(0, 1, 0, 0, 0, 0, 0, "读取输入图像", Path(self.input_path).name, 5)

            input_image = load_image(self.input_path)
            label_image = None
            label_path = self.label_path or ""
            if label_path:
                try:
                    label_image = load_image(label_path)
                    self.log_message.emit(f"[OK] 已加载真实红外标签: {label_path}")
                except Exception as exc:
                    label_path = ""
                    self.log_message.emit(f"[WARN] 标签图读取失败，跳过指标和标签对比: {exc}")
            else:
                self.log_message.emit("[INFO] 未选择真实红外标签，后续跳过质量指标计算")

            self.preview_ready.emit(input_image, None, label_image, None)
            if self.cancel_requested:
                self.log_message.emit("[CANCEL] 用户取消任务，未开始模型推理")
                self.result_ready.emit({"record": record, "run_dir": run_dir, "canceled": True})
                return

            self._emit_progress(0, 1, 0, 0, 0, time.time() - start_clock, 0, "模型推理中", Path(self.input_path).name, 45)
            self.log_message.emit("[INFO] 模型加载/复用中，开始推理")
            generated_image, elapsed = _call_generation_algorithm(
                algorithm_id,
                input_image,
                self.options["device_preference"],
            )
            self.log_message.emit(f"[OK] 单张推理完成 | 单张推理耗时: {elapsed:.3f}s")

            stem = _result_stem(self.input_path, algorithm_key, self.options.get("filename_prefix", ""))
            output_path = save_image_to_path(
                generated_image,
                os.path.join(generated_dir, f"{stem}_ir.png"),
                overwrite=True,
            )
            self.log_message.emit(f"[OK] 结果保存路径: {output_path}")

            compare_path = ""
            compare_image = None
            if self.options.get("save_compare", True):
                compare_image = create_comparison_image(input_image, generated_image, label_image)
                compare_path = save_image_to_path(
                    compare_image,
                    os.path.join(compare_dir, f"{stem}_compare.png"),
                    overwrite=True,
                )
                self.log_message.emit(f"[OK] 对比可视化图保存路径: {compare_path}")

            metrics = {"psnr": None, "ssim": None, "l1": None, "lipips": None}
            if self.options.get("compute_metrics", True) and label_image is not None:
                metrics = calculate_metrics(generated_image, label_image)
                if all(metrics.get(key) is None for key in ("psnr", "ssim", "l1", "lipips")):
                    self.log_message.emit("[WARN] 指标计算失败或返回空值，已跳过写入")
                else:
                    self.log_message.emit(
                        "[OK] 指标计算结果: "
                        f"PSNR={metrics.get('psnr')} SSIM={metrics.get('ssim')} "
                        f"L1={metrics.get('l1')} LIPIPS={metrics.get('lipips')}"
                    )
                    if metrics.get("lipips") is None:
                        reason = get_lipips_error() or "LPIPS runtime returned no value"
                        self.log_message.emit(f"[WARN] LIPIPS 无法计算: {reason}")
            elif not self.options.get("compute_metrics", True):
                self.log_message.emit("[INFO] 已关闭质量指标计算")
            else:
                self.log_message.emit("[INFO] 未选择真实红外标签，跳过质量指标计算")

            record.update(
                {
                    "label_path": label_path,
                    "output_path": output_path,
                    "compare_path": compare_path,
                    "status": "success",
                    "error_message": "",
                    "inference_time": f"{elapsed:.4f}",
                    "psnr": _metric_text(metrics.get("psnr")),
                    "ssim": _metric_text(metrics.get("ssim")),
                    "l1": _metric_text(metrics.get("l1")),
                    "lipips": _metric_text(metrics.get("lipips")),
                }
            )
            self.preview_ready.emit(input_image, generated_image, label_image, compare_image)

            params_path = os.path.join(reports_dir, "params.txt")
            _write_params_file(params, params_path)
            report_txt = ""
            if self.options.get("export_report", True):
                report_txt = os.path.join(reports_dir, "single_generation_report.txt")
                try:
                    _write_single_txt_report(record, params, report_txt)
                    self.log_message.emit(f"[OK] TXT 报告已导出: {report_txt}")
                except Exception as exc:
                    self.log_message.emit(f"[ERROR] 报告写入失败，生成结果已保留: {exc}")
                self._maybe_export_docx(reports_dir, record)
            else:
                self.log_message.emit("[INFO] 已关闭报告导出，仅保存生成图像和参数记录")

            end_readable = time.strftime("%Y-%m-%d %H:%M:%S")
            total_elapsed = time.time() - start_clock
            summary = {
                "start_time": start_readable,
                "end_time": end_readable,
                "total_elapsed": f"{total_elapsed:.3f}",
                "params_path": params_path,
            }
            self._emit_progress(1, 1, 1, 0, elapsed, total_elapsed, total_elapsed, "已完成", Path(self.input_path).name, 100)
            self.log_message.emit(f"[DONE] 单张生成完成 | 总耗时: {total_elapsed:.3f}s | 输出目录: {run_dir}")
            self.result_ready.emit(
                {
                    "record": record,
                    "summary": summary,
                    "run_dir": run_dir,
                    "params_path": params_path,
                    "report_txt": report_txt,
                    "canceled": False,
                }
            )
        except Exception as exc:
            record["status"] = "failed"
            record["error_message"] = str(exc)
            self.error_occurred.emit(str(exc))

    def _empty_record(self):
        return {
            "image_name": Path(self.input_path).name,
            "input_path": self.input_path,
            "label_path": self.label_path or "",
            "output_path": "",
            "compare_path": "",
            "algorithm": self.options["algorithm_key"],
            "device": self.options["device"],
            "status": "failed",
            "error_message": "",
            "inference_time": "",
            "psnr": "",
            "ssim": "",
            "l1": "",
            "lipips": "",
        }

    def _params(self, run_dir):
        return {
            "algorithm": self.options["algorithm_key"],
            "device": self.options["device"],
            "input mode": "single",
            "input path": self.input_path,
            "output directory": run_dir,
            "label path": self.label_path or "",
            "label directory": "",
            "save_compare": self.options.get("save_compare", True),
            "export_report": self.options.get("export_report", True),
            "export_docx": self.options.get("export_docx", False),
            "compute_metrics": self.options.get("compute_metrics", True),
            "filename_prefix": self.options.get("filename_prefix", ""),
            "image_formats": ", ".join(IMAGE_EXTENSIONS),
        }

    def _maybe_export_docx(self, reports_dir, record):
        if not self.options.get("export_docx", False):
            return
        docx_path = os.path.join(reports_dir, "single_generation_report.docx")
        try:
            _try_write_docx_report(
                "单张红外图像生成报告",
                [_record_report_line(key, value) for key, value in record.items()],
                docx_path,
            )
            self.log_message.emit(f"[OK] DOCX 报告已导出: {docx_path}")
        except ImportError:
            self.log_message.emit("[WARN] 未安装 python-docx，跳过 docx 报告导出")
        except Exception as exc:
            self.log_message.emit(f"[WARN] DOCX 报告导出失败，已跳过: {exc}")

    def _emit_progress(self, completed, total, success, failed, current_elapsed, total_elapsed, avg_elapsed, status, current_file, percent):
        self.progress_changed.emit(
            {
                "mode": "single",
                "completed": completed,
                "total": total,
                "success": success,
                "failed": failed,
                "current_elapsed": current_elapsed,
                "total_elapsed": total_elapsed,
                "avg_elapsed": avg_elapsed,
                "status": status,
                "current_file": current_file,
                "percent": percent,
            }
        )


class BatchGenerationWorker(QThread):
    """批量红外图像生成工作线程。"""

    log_message = pyqtSignal(str)
    progress_changed = pyqtSignal(object)
    preview_ready = pyqtSignal(object, object, object, object)
    result_ready = pyqtSignal(object)
    error_occurred = pyqtSignal(str)

    def __init__(self, input_files, input_dir, label_dir, options):
        super().__init__()
        self.input_files = input_files
        self.input_dir = input_dir
        self.label_dir = label_dir
        self.options = options
        self.cancel_requested = False

    def request_cancel(self):
        self.cancel_requested = True

    def run(self):
        start_readable = time.strftime("%Y-%m-%d %H:%M:%S")
        start_clock = time.time()
        algorithm_key = self.options["algorithm_key"]
        output_root = self.options["output_dir"]
        batch_root = os.path.join(output_root, "批次文件夹")
        batch_folder_name = Path(self.input_dir).name
        run_dir = _unique_dir(batch_root, batch_folder_name)
        generated_dir = os.path.join(run_dir, "generated")
        compare_dir = os.path.join(run_dir, "compare")
        reports_dir = os.path.join(run_dir, "reports")

        try:
            for folder in (generated_dir, compare_dir, reports_dir):
                os.makedirs(folder, exist_ok=True)
        except Exception as exc:
            self.error_occurred.emit(f"输出目录无法创建: {exc}")
            return

        total = len(self.input_files)
        records = []
        success = 0
        failed = 0
        canceled = False
        params = self._params(run_dir)

        self.log_message.emit("[RUN] 启动批量生成任务")
        self.log_message.emit("[INFO] 当前处理模式: 批量")
        self.log_message.emit(f"[INFO] 输入路径: {self.input_dir}")
        self.log_message.emit(f"[INFO] 标签路径: {self.label_dir or '未选择'}")
        self.log_message.emit(f"[INFO] 输出路径: {run_dir}")
        self.log_message.emit(f"[INFO] 使用算法: {algorithm_key} | 使用设备: {self.options['device']}")
        self.log_message.emit(f"[INFO] 待处理图像数量: {total}")

        for index, image_path in enumerate(self.input_files, 1):
            if self.cancel_requested:
                canceled = True
                self.log_message.emit("[CANCEL] 用户取消任务，停止后续图片处理")
                break

            image_start = time.time()
            current_file = Path(image_path).name
            record = self._empty_record(image_path)
            self._emit_progress(index - 1, total, success, failed, 0, time.time() - start_clock, 0, "处理中", current_file)
            self.log_message.emit(f"[RUN] ({index}/{total}) 当前处理图片名: {current_file}")

            try:
                input_image = load_image(image_path)
                label_path, label_image = self._load_label_for(image_path)
                self.preview_ready.emit(input_image, None, label_image, None)

                generated_image, elapsed = _call_generation_algorithm(
                    self.options["algorithm_id"],
                    input_image,
                    self.options["device_preference"],
                )
                stem = _result_stem(image_path, algorithm_key, self.options.get("filename_prefix", ""))
                output_path = save_image_to_path(
                    generated_image,
                    os.path.join(generated_dir, f"{stem}_ir.png"),
                    overwrite=True,
                )

                compare_path = ""
                compare_image = None
                if self.options.get("save_compare", True):
                    compare_image = create_comparison_image(input_image, generated_image, label_image)
                    compare_path = save_image_to_path(
                        compare_image,
                        os.path.join(compare_dir, f"{stem}_compare.png"),
                        overwrite=True,
                    )

                metrics = {"psnr": None, "ssim": None, "l1": None, "lipips": None}
                if self.options.get("compute_metrics", True) and label_image is not None:
                    metrics = calculate_metrics(generated_image, label_image)
                    if all(metrics.get(key) is None for key in ("psnr", "ssim", "l1", "lipips")):
                        self.log_message.emit("[WARN] 指标计算失败或返回空值，已跳过写入")
                    else:
                        self.log_message.emit(
                            f"[OK] 指标计算结果: PSNR={metrics.get('psnr')} SSIM={metrics.get('ssim')} "
                            f"L1={metrics.get('l1')} LIPIPS={metrics.get('lipips')}"
                        )
                        if metrics.get("lipips") is None:
                            reason = get_lipips_error() or "LPIPS runtime returned no value"
                            self.log_message.emit(f"[WARN] LIPIPS 无法计算: {reason}")
                elif not self.options.get("compute_metrics", True):
                    self.log_message.emit("[INFO] 已关闭质量指标计算")
                else:
                    self.log_message.emit("[INFO] 未选择真实红外标签，跳过质量指标计算")

                record.update(
                    {
                        "label_path": label_path,
                        "output_path": output_path,
                        "compare_path": compare_path,
                        "status": "success",
                        "error_message": "",
                        "inference_time": f"{elapsed:.4f}",
                        "psnr": _metric_text(metrics.get("psnr")),
                        "ssim": _metric_text(metrics.get("ssim")),
                        "l1": _metric_text(metrics.get("l1")),
                        "lipips": _metric_text(metrics.get("lipips")),
                    }
                )
                success += 1
                self.preview_ready.emit(input_image, generated_image, label_image, compare_image)
                self.log_message.emit(f"[OK] 结果保存路径: {output_path}")
                if compare_path:
                    self.log_message.emit(f"[OK] 对比可视化图保存路径: {compare_path}")
                self.log_message.emit(f"[OK] 单张推理耗时: {elapsed:.3f}s")
            except Exception as exc:
                failed += 1
                record["status"] = "failed"
                record["error_message"] = str(exc)
                self.log_message.emit(f"[ERROR] 失败图片: {current_file} | 失败原因: {exc}")

            records.append(record)
            completed = success + failed
            total_elapsed = time.time() - start_clock
            current_elapsed = time.time() - image_start
            avg_elapsed = total_elapsed / completed if completed else 0
            self._emit_progress(completed, total, success, failed, current_elapsed, total_elapsed, avg_elapsed, "处理中", current_file)

        completed = success + failed
        total_elapsed = time.time() - start_clock
        avg_elapsed = total_elapsed / completed if completed else 0
        task_status = "canceled" if canceled else "completed"
        summary = {
            "input_dir": self.input_dir,
            "output_dir": run_dir,
            "label_dir": self.label_dir or "",
            "algorithm": algorithm_key,
            "device": self.options["device"],
            "start_time": start_readable,
            "end_time": time.strftime("%Y-%m-%d %H:%M:%S"),
            "task_status": task_status,
            "total": total,
            "completed": completed,
            "success": success,
            "failed": failed,
            "total_elapsed": total_elapsed,
            "avg_elapsed": avg_elapsed,
        }

        report_txt = os.path.join(reports_dir, "batch_generation_results.txt")
        try:
            _write_batch_results_txt(summary, records, params, report_txt)
            self.log_message.emit(f"[OK] 批量 TXT 结果总表已导出: {report_txt}")
        except Exception as exc:
            self.log_message.emit(f"[ERROR] 报告写入失败，已保留生成结果: {exc}")

        self._emit_progress(completed, total, success, failed, 0, total_elapsed, avg_elapsed, "已取消" if canceled else "已完成", "")
        self.log_message.emit(
            f"[DONE] 任务完成摘要 | 状态={task_status} | 成功={success} | 失败={failed} "
            f"| 已处理={completed}/{total} | 总耗时={total_elapsed:.3f}s"
        )
        self.result_ready.emit(
            {
                "records": records,
                "summary": summary,
                "run_dir": run_dir,
                "generated_dir": generated_dir,
                "compare_dir": compare_dir,
                "reports_dir": reports_dir,
                "report_txt": report_txt,
                "canceled": canceled,
            }
        )

    def _load_label_for(self, image_path):
        if not self.label_dir:
            return "", None
        try:
            label_path = find_matching_label(image_path, self.label_dir) or ""
            if not label_path:
                self.log_message.emit(f"[INFO] 找不到对应标签图，跳过指标: {Path(image_path).name}")
                return "", None
            label_image = load_image(label_path)
            self.log_message.emit(f"[OK] 匹配真实红外标签: {Path(label_path).name}")
            return label_path, label_image
        except Exception as exc:
            self.log_message.emit(f"[WARN] 标签图读取失败，跳过指标: {exc}")
            return "", None

    def _empty_record(self, image_path):
        return {
            "image_name": Path(image_path).name,
            "input_path": image_path,
            "label_path": "",
            "output_path": "",
            "compare_path": "",
            "algorithm": self.options["algorithm_key"],
            "device": self.options["device"],
            "status": "failed",
            "error_message": "",
            "inference_time": "",
            "psnr": "",
            "ssim": "",
            "l1": "",
            "lipips": "",
        }

    def _params(self, run_dir):
        return {
            "algorithm": self.options["algorithm_key"],
            "device": self.options["device"],
            "input mode": "batch",
            "input path": self.input_dir,
            "output directory": run_dir,
            "label path": "",
            "label directory": self.label_dir or "",
            "save_compare": self.options.get("save_compare", True),
            "export_report": True,
            "export_docx": self.options.get("export_docx", False),
            "compute_metrics": self.options.get("compute_metrics", True),
            "filename_prefix": self.options.get("filename_prefix", ""),
            "image_formats": ", ".join(IMAGE_EXTENSIONS),
        }

    def _emit_progress(self, completed, total, success, failed, current_elapsed, total_elapsed, avg_elapsed, status, current_file):
        percent = int(completed / total * 100) if total else 0
        self.progress_changed.emit(
            {
                "mode": "batch",
                "completed": completed,
                "total": total,
                "success": success,
                "failed": failed,
                "current_elapsed": current_elapsed,
                "total_elapsed": total_elapsed,
                "avg_elapsed": avg_elapsed,
                "status": status,
                "current_file": current_file,
                "percent": percent,
            }
        )


class GenerationTab(QWidget):
    """可见光到红外生成完整工作流页面。"""

    ALGORITHM_HINTS = {
        1: "DFSMamba 为默认算法，支持 CPU/CUDA 设备选择，CUDA 不可用时后端可回退 CPU。",
        2: "PhysMamba 为默认算法，支持 CPU/CUDA 设备选择，CUDA 不可用时后端可回退 CPU。",
    }

    def __init__(self, log_signal=None):
        super().__init__()
        self.log_signal = log_signal
        self.input_image = None
        self.generated_image = None
        self.gt_image = None
        self.compare_image = None
        self.gen_worker = None
        self.metric_labels = {}
        self.progress_labels = {}
        self._init_ui()

    def _init_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(18, 14, 18, 16)
        main_layout.setSpacing(14)
        main_layout.addWidget(self._create_header())

        workspace = QHBoxLayout()
        workspace.setSpacing(14)
        workspace.addWidget(self._create_workflow_panel(), 0)
        workspace.addWidget(self._create_preview_panel(), 1)
        workspace.addWidget(self._create_status_panel(), 0)
        main_layout.addLayout(workspace, 1)
        main_layout.addWidget(self._create_log_panel(), 0)

        os.makedirs(DEFAULT_OUTPUT_DIR, exist_ok=True)
        self.lbl_output_dir.setText(DEFAULT_OUTPUT_DIR)
        self.lbl_output_dir.setToolTip(DEFAULT_OUTPUT_DIR)
        self.combo_algorithm.setCurrentIndex(1)
        self._on_algorithm_changed()
        self._on_device_changed()
        self._on_mode_changed()
        self._reset_progress()

    def _create_header(self):
        header = QFrame()
        header.setObjectName("pageHeader")
        layout = QHBoxLayout(header)
        layout.setContentsMargins(18, 14, 18, 14)
        layout.setSpacing(16)

        title_area = QVBoxLayout()
        title_area.setSpacing(4)
        title = QLabel("可见光到红外图像生成工作流")
        title.setObjectName("pageTitle")
        subtitle = QLabel("Import Data  |  Configure  |  Run  |  Preview  |  Archive Results")
        subtitle.setObjectName("pageSubtitle")
        title_area.addWidget(title)
        title_area.addWidget(subtitle)
        layout.addLayout(title_area, 1)

        status_card = QFrame()
        status_card.setObjectName("statusBadge")
        status_layout = QGridLayout(status_card)
        status_layout.setContentsMargins(14, 8, 14, 8)
        status_layout.setHorizontalSpacing(12)
        status_layout.setVerticalSpacing(4)
        self.lbl_algorithm_status = self._status_value("PhysMamba")
        self.lbl_generation_status = self._status_value("待运行")
        self.lbl_device_status = self._status_value("CPU")
        self.lbl_current_file = self._status_value("-")
        self.lbl_current_file.setMinimumWidth(180)
        for row, (name, value) in enumerate(
            [
                ("当前算法", self.lbl_algorithm_status),
                ("运行状态", self.lbl_generation_status),
                ("运行设备", self.lbl_device_status),
                ("当前图像", self.lbl_current_file),
            ]
        ):
            label = QLabel(name)
            label.setObjectName("statusLabel")
            status_layout.addWidget(label, row, 0)
            status_layout.addWidget(value, row, 1)
        layout.addWidget(status_card, 0)
        return header

    def _create_workflow_panel(self):
        panel = QFrame()
        panel.setObjectName("configPanel")
        panel.setMinimumWidth(400)
        panel.setMaximumWidth(430)
        outer_layout = QVBoxLayout(panel)
        outer_layout.setContentsMargins(16, 16, 16, 16)
        outer_layout.setSpacing(12)

        scroll_area = QScrollArea()
        scroll_area.setObjectName("workflowScrollArea")
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        scroll_content = QWidget()
        scroll_content.setObjectName("workflowScrollContent")
        layout = QVBoxLayout(scroll_content)
        layout.setContentsMargins(0, 0, 6, 0)
        layout.setSpacing(10)

        layout.addWidget(self._section_title("数据导入"))
        mode_row = QHBoxLayout()
        self.radio_single = QRadioButton("单张处理")
        self.radio_batch = QRadioButton("批量处理")
        self.radio_single.setChecked(True)
        self.radio_single.toggled.connect(self._on_mode_changed)
        mode_row.addWidget(self.radio_single)
        mode_row.addWidget(self.radio_batch)
        mode_row.addStretch(1)
        layout.addLayout(mode_row)

        self.single_input_field, self.lbl_input_path, self.btn_select_input = self._add_file_picker(
            layout, "单张可见光图像", "请选择输入图像...", self._on_select_input
        )
        self.batch_input_field, self.lbl_input_folder, self.btn_select_input_folder = self._add_file_picker(
            layout, "批量输入文件夹", "请选择输入文件夹...", self._on_select_input_folder, folder=True
        )
        self.single_label_field, self.lbl_gt_path, self.btn_select_gt = self._add_file_picker(
            layout, "单张真实红外标签（可选）", "用于指标和对比图...", self._on_select_gt
        )
        self.batch_label_field, self.lbl_gt_folder, self.btn_select_gt_folder = self._add_file_picker(
            layout, "批量标签文件夹（可选）", "按同名文件匹配标签...", self._on_select_gt_folder, folder=True
        )

        layout.addWidget(self._divider())
        layout.addWidget(self._section_title("参数设置"))

        algo_row = QHBoxLayout()
        algo_col = QVBoxLayout()
        algo_col.addWidget(self._field_label("生成算法"))
        self.combo_algorithm = QComboBox()
        self.combo_algorithm.addItem("算法一：DFSMamba", 1)
        self.combo_algorithm.addItem("算法二：PhysMamba", 2)
        self.combo_algorithm.currentIndexChanged.connect(self._on_algorithm_changed)
        algo_col.addWidget(self.combo_algorithm)
        device_col = QVBoxLayout()
        device_col.addWidget(self._field_label("运行设备"))
        self.combo_device = QComboBox()
        self.combo_device.addItem("CPU", "cpu")
        self.combo_device.addItem("CUDA", "cuda")
        self.combo_device.currentIndexChanged.connect(self._on_device_changed)
        device_col.addWidget(self.combo_device)
        algo_row.addLayout(algo_col, 2)
        algo_row.addLayout(device_col, 1)
        layout.addLayout(algo_row)

        self.algorithm_hint = QLabel(self.ALGORITHM_HINTS[2])
        self.algorithm_hint.setObjectName("sectionHint")
        self.algorithm_hint.setWordWrap(True)
        layout.addWidget(self.algorithm_hint)

        self.output_dir_field, self.lbl_output_dir, self.btn_select_output_dir = self._add_file_picker(
            layout, "输出目录", "请选择输出目录...", self._on_select_output_dir, folder=True
        )
        prefix_row = QHBoxLayout()
        prefix_row.addWidget(self._field_label("输出文件名前缀"))
        self.edit_filename_prefix = QLineEdit()
        self.edit_filename_prefix.setPlaceholderText("可选，例如 exp01")
        prefix_row.addWidget(self.edit_filename_prefix, 1)
        layout.addLayout(prefix_row)

        self.chk_compute_metrics = QCheckBox("计算质量指标（有标签时）")
        self.chk_compute_metrics.setChecked(True)
        self.chk_save_compare = QCheckBox("保存对比可视化图")
        self.chk_save_compare.setChecked(True)
        self.chk_export_report = QCheckBox("导出 TXT 报告（批量写入一张总表）")
        self.chk_export_report.setChecked(True)
        self.chk_export_docx = QCheckBox("导出 DOCX 报告（可选）")
        for checkbox in (
            self.chk_compute_metrics,
            self.chk_save_compare,
            self.chk_export_report,
            self.chk_export_docx,
        ):
            layout.addWidget(checkbox)

        layout.addStretch(1)
        scroll_area.setWidget(scroll_content)
        outer_layout.addWidget(scroll_area, 1)

        footer = QFrame()
        footer.setObjectName("workflowFooter")
        footer_layout = QVBoxLayout(footer)
        footer_layout.setContentsMargins(0, 0, 0, 0)
        footer_layout.setSpacing(8)

        footer_layout.addWidget(self._divider())
        footer_layout.addWidget(self._section_title("运行算法"))
        self.btn_start_single = QPushButton("开始生成单张")
        self.btn_start_single.setObjectName("btnGenerate")
        self.btn_start_single.clicked.connect(self._on_generate_single)
        footer_layout.addWidget(self.btn_start_single)
        self.btn_start_batch = QPushButton("开始批量生成")
        self.btn_start_batch.setObjectName("btnGenerate")
        self.btn_start_batch.clicked.connect(self._on_generate_batch)
        footer_layout.addWidget(self.btn_start_batch)
        self.btn_cancel = QPushButton("取消当前任务")
        self.btn_cancel.setEnabled(False)
        self.btn_cancel.clicked.connect(self._on_cancel_task)
        footer_layout.addWidget(self.btn_cancel)
        self.btn_open_dir = QPushButton("打开输出目录")
        self.btn_open_dir.setObjectName("btnOpenDir")
        self.btn_open_dir.clicked.connect(self._on_open_save_dir)
        footer_layout.addWidget(self.btn_open_dir)
        outer_layout.addWidget(footer, 0)
        return panel

    def _create_preview_panel(self):
        panel = QFrame()
        panel.setObjectName("panelCard")
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)
        title = QLabel("结果预览")
        title.setObjectName("sectionTitle")
        hint = QLabel("批量处理时显示最近一次处理成功的图像结果。")
        hint.setObjectName("sectionHint")
        layout.addWidget(title)
        layout.addWidget(hint)
        grid = QGridLayout()
        grid.setHorizontalSpacing(12)
        grid.setVerticalSpacing(12)
        input_card, self.lbl_input_image = self._create_image_card("输入可见光图像", "Input", "等待导入")
        output_card, self.lbl_output_image = self._create_image_card("生成红外图像", "Generated IR", "等待生成")
        gt_card, self.lbl_gt_image = self._create_image_card("真实红外标签", "Ground Truth", "未选择标签")
        compare_card, self.lbl_compare_image = self._create_image_card("对比可视化图", "Comparison", "等待生成对比图")
        grid.addWidget(input_card, 0, 0)
        grid.addWidget(output_card, 0, 1)
        grid.addWidget(gt_card, 1, 0)
        grid.addWidget(compare_card, 1, 1)
        layout.addLayout(grid, 1)
        return panel

    def _create_status_panel(self):
        panel = QFrame()
        panel.setObjectName("statsPanel")
        panel.setMinimumWidth(330)
        panel.setMaximumWidth(370)
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)
        layout.addWidget(self._section_title("进度显示"))
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        layout.addWidget(self.progress_bar)
        progress_grid = QGridLayout()
        for index, (key, name, default) in enumerate(
            [
                ("completed", "已完成/总数", "0/0"),
                ("success", "成功数量", "0"),
                ("failed", "失败数量", "0"),
                ("current_elapsed", "当前耗时", "0.000 s"),
                ("total_elapsed", "总耗时", "0.000 s"),
                ("avg_elapsed", "平均单张耗时", "0.000 s"),
            ]
        ):
            progress_grid.addWidget(self._create_value_card(self.progress_labels, key, name, default), index // 2, index % 2)
        layout.addLayout(progress_grid)

        layout.addWidget(self._divider())
        layout.addWidget(self._section_title("质量指标"))
        metric_grid = QGridLayout()
        for index, (key, name, default) in enumerate(
            [
                ("psnr", "PSNR", "未提供标签"),
                ("ssim", "SSIM", "未提供标签"),
                ("l1", "L1", "未提供标签"),
                ("lipips", "LIPIPS", "未提供标签"),
            ]
        ):
            metric_grid.addWidget(self._create_value_card(self.metric_labels, key, name, default), index // 2, index % 2)
        layout.addLayout(metric_grid)
        layout.addWidget(self._field_label("最近保存路径"))
        self.metric_labels["save_path"] = QLabel("未保存")
        self.metric_labels["save_path"].setObjectName("pathValue")
        self.metric_labels["save_path"].setWordWrap(True)
        layout.addWidget(self.metric_labels["save_path"])
        layout.addWidget(self._field_label("最近报告路径"))
        self.metric_labels["report_path"] = QLabel("未导出")
        self.metric_labels["report_path"].setObjectName("pathValue")
        self.metric_labels["report_path"].setWordWrap(True)
        layout.addWidget(self.metric_labels["report_path"])
        layout.addStretch(1)
        return panel

    def _create_log_panel(self):
        panel = QFrame()
        panel.setObjectName("logCard")
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)
        title = QLabel("运行日志")
        title.setObjectName("sectionTitle")
        layout.addWidget(title)
        self.log_text = QTextEdit()
        self.log_text.setObjectName("logText")
        self.log_text.setReadOnly(True)
        self.log_text.setFont(QFont("Consolas", 9))
        self.log_text.setMaximumHeight(118)
        layout.addWidget(self.log_text)
        return panel

    def _status_value(self, text):
        label = QLabel(text)
        label.setObjectName("statusValue")
        return label

    def _section_title(self, text):
        label = QLabel(text)
        label.setObjectName("sectionTitle")
        return label

    def _field_label(self, text):
        label = QLabel(text)
        label.setObjectName("fieldLabel")
        return label

    def _divider(self):
        divider = QFrame()
        divider.setObjectName("subtleDivider")
        divider.setFrameShape(QFrame.HLine)
        divider.setMaximumHeight(1)
        return divider

    def _add_file_picker(self, parent_layout, title, placeholder, slot, folder=False):
        field = QFrame()
        field.setObjectName("workflowField")
        field_layout = QVBoxLayout(field)
        field_layout.setContentsMargins(0, 0, 0, 0)
        field_layout.setSpacing(6)

        field_layout.addWidget(self._field_label(title))
        row = QHBoxLayout()
        row.setSpacing(8)
        line_edit = QLineEdit()
        line_edit.setReadOnly(True)
        line_edit.setPlaceholderText(placeholder)
        button = QPushButton("选择文件夹" if folder else "选择")
        button.clicked.connect(slot)
        row.addWidget(line_edit, 1)
        row.addWidget(button, 0)
        field_layout.addLayout(row)
        parent_layout.addWidget(field)
        return field, line_edit, button

    def _create_image_card(self, title, subtitle, placeholder):
        card = QFrame()
        card.setObjectName("imageCard")
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(12, 12, 12, 12)
        card_layout.setSpacing(8)
        title_label = QLabel(title)
        title_label.setObjectName("imageTitle")
        subtitle_label = QLabel(subtitle)
        subtitle_label.setObjectName("imageHint")
        card_layout.addWidget(title_label)
        card_layout.addWidget(subtitle_label)
        canvas = QFrame()
        canvas.setObjectName("imageCanvasFrame")
        canvas_layout = QVBoxLayout(canvas)
        canvas_layout.setContentsMargins(0, 0, 0, 0)
        image_label = QLabel(placeholder)
        image_label.setObjectName("imageCanvas")
        image_label.setAlignment(Qt.AlignCenter)
        image_label.setMinimumSize(260, 210)
        image_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        image_label.setWordWrap(True)
        canvas_layout.addWidget(image_label)
        card_layout.addWidget(canvas, 1)
        return card, image_label

    def _create_value_card(self, target, key, title, default_value):
        card = QFrame()
        card.setObjectName("summaryCard")
        card.setProperty("accent", "generation")
        layout = QVBoxLayout(card)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(4)
        name_label = QLabel(title)
        name_label.setObjectName("summaryName")
        value_label = QLabel(default_value)
        value_label.setObjectName("summaryValue")
        value_label.setWordWrap(True)
        layout.addWidget(name_label)
        layout.addWidget(value_label)
        target[key] = value_label
        return card

    def _on_algorithm_changed(self):
        if not hasattr(self, "combo_algorithm"):
            return
        algorithm_id = self.combo_algorithm.currentData()
        hint = self.ALGORITHM_HINTS.get(algorithm_id, "")
        if algorithm_id == 1 and self._current_device_preference() == "cpu":
            hint = "DFSMamba 为默认算法，支持 CPU/CUDA 设备选择，CUDA 不可用时后端可回退 CPU。"
        self.lbl_algorithm_status.setText(_algorithm_key(algorithm_id))
        if hasattr(self, "algorithm_hint"):
            self.algorithm_hint.setText(hint)

    def _on_device_changed(self):
        if not hasattr(self, "combo_device"):
            return
        self.lbl_device_status.setText(self._current_device_label())
        self._on_algorithm_changed()

    def _on_mode_changed(self):
        if not hasattr(self, "btn_start_single"):
            return
        single_mode = self.radio_single.isChecked()
        running = self._is_generation_running()

        for widget in (self.single_input_field, self.single_label_field):
            widget.setVisible(single_mode)
        for widget in (self.batch_input_field, self.batch_label_field):
            widget.setVisible(not single_mode)

        self.chk_export_docx.setVisible(single_mode)
        self.btn_start_single.setVisible(single_mode)
        self.btn_start_batch.setVisible(not single_mode)
        self.btn_start_single.setEnabled(single_mode and not running)
        self.btn_start_batch.setEnabled((not single_mode) and not running)
        self.lbl_generation_status.setText("单张模式待运行" if single_mode else "批量模式待运行")

    def _current_device_preference(self):
        if not hasattr(self, "combo_device"):
            return "cpu"
        return self.combo_device.currentData() or "cpu"

    def _current_device_label(self):
        if not hasattr(self, "combo_device"):
            return "CPU"
        return self.combo_device.currentText() or "CPU"

    def _current_options(self):
        algorithm_id = self.combo_algorithm.currentData()
        return {
            "algorithm_id": algorithm_id,
            "algorithm_key": _algorithm_key(algorithm_id),
            "device_preference": self._current_device_preference(),
            "device": self._current_device_label(),
            "output_dir": self.lbl_output_dir.text().strip() or DEFAULT_OUTPUT_DIR,
            "compute_metrics": self.chk_compute_metrics.isChecked(),
            "save_compare": self.chk_save_compare.isChecked(),
            "export_report": self.chk_export_report.isChecked(),
            "export_docx": self.chk_export_docx.isChecked(),
            "filename_prefix": self.edit_filename_prefix.text().strip(),
        }

    def _on_select_input(self):
        if self._warn_if_running():
            return
        path, _ = QFileDialog.getOpenFileName(self, "选择可见光图像", "", IMAGE_FILTER)
        if not path:
            return
        try:
            self.input_image = load_image(path)
            self.generated_image = None
            self.compare_image = None
            self.lbl_input_path.setText(path)
            self.lbl_input_path.setToolTip(path)
            self.lbl_output_image.setText("等待生成")
            self.lbl_compare_image.setText("等待生成对比图")
            show_image_on_label(self.lbl_input_image, self.input_image)
            self._log(f"[OK] 已加载输入图像: {os.path.basename(path)}")
        except Exception as exc:
            QMessageBox.warning(self, "加载失败", str(exc))
            self._log(f"[ERROR] 加载图像失败: {exc}")

    def _on_select_input_folder(self):
        if self._warn_if_running():
            return
        path = QFileDialog.getExistingDirectory(self, "选择批量输入文件夹", "")
        if not path:
            return
        self.lbl_input_folder.setText(path)
        self.lbl_input_folder.setToolTip(path)
        try:
            files = scan_image_files(path)
            self._log(f"[OK] 已选择输入文件夹: {path} | 支持格式图像 {len(files)} 张")
        except Exception as exc:
            QMessageBox.warning(self, "输入文件夹无效", str(exc))
            self._log(f"[ERROR] 输入文件夹扫描失败: {exc}")

    def _on_select_gt(self):
        if self._warn_if_running():
            return
        path, _ = QFileDialog.getOpenFileName(self, "选择真实红外标签图像", "", IMAGE_FILTER)
        if not path:
            return
        try:
            self.gt_image = load_image(path)
            self.lbl_gt_path.setText(path)
            self.lbl_gt_path.setToolTip(path)
            show_image_on_label(self.lbl_gt_image, self.gt_image)
            self._log(f"[OK] 已加载标签图像: {os.path.basename(path)}")
        except Exception as exc:
            QMessageBox.warning(self, "加载失败", str(exc))
            self._log(f"[ERROR] 标签图像加载失败: {exc}")

    def _on_select_gt_folder(self):
        if self._warn_if_running():
            return
        path = QFileDialog.getExistingDirectory(self, "选择真实红外标签文件夹", "")
        if not path:
            return
        self.lbl_gt_folder.setText(path)
        self.lbl_gt_folder.setToolTip(path)
        self._log(f"[OK] 已选择标签文件夹: {path}")

    def _on_select_output_dir(self):
        if self._warn_if_running():
            return
        path = QFileDialog.getExistingDirectory(self, "选择输出目录", self.lbl_output_dir.text() or DEFAULT_OUTPUT_DIR)
        if not path:
            return
        self.lbl_output_dir.setText(path)
        self.lbl_output_dir.setToolTip(path)
        self._log(f"[OK] 已选择输出目录: {path}")

    def _validate_algorithm_device(self):
        algorithm_id = self.combo_algorithm.currentData()
        device = self._current_device_preference()
        if algorithm_id == 1 and device == "cpu":
            QMessageBox.warning(self, "设备不支持", "DFSMamba 当前暂未适配 CPU 推理，请切换 CUDA 或选择 PhysMamba。")
            self._log("[ERROR] DFSMamba 当前暂未适配 CPU 推理")
            return False
        if device == "cuda" and not self._cuda_available():
            if algorithm_id == 1:
                QMessageBox.warning(self, "CUDA 不可用", "当前环境未检测到可用 CUDA，无法运行 DFSMamba。")
                self._log("[ERROR] CUDA 不可用，DFSMamba 无法运行")
                return False
            QMessageBox.information(self, "CUDA 不可用", "当前环境未检测到可用 CUDA，PhysMamba 后端将回退到 CPU。")
            self._log("[WARN] CUDA 不可用，PhysMamba 将由后端回退到 CPU")
        return True

    def _cuda_available(self):
        try:
            import torch

            return bool(torch.cuda.is_available())
        except Exception:
            return False

    def _ensure_output_dir(self):
        output_dir = self.lbl_output_dir.text().strip() or DEFAULT_OUTPUT_DIR
        try:
            os.makedirs(output_dir, exist_ok=True)
        except Exception as exc:
            QMessageBox.warning(self, "输出目录错误", f"输出目录无法创建:\n{exc}")
            self._log(f"[ERROR] 输出目录无法创建: {exc}")
            return None
        return output_dir

    def _on_generate_single(self):
        if self._is_generation_running():
            return
        input_path = self.lbl_input_path.text().strip()
        if not input_path:
            QMessageBox.warning(self, "缺少输入图像", "请先选择一张可见光图像。")
            self._log("[ERROR] 没有选择输入图像")
            return
        if not os.path.isfile(input_path):
            QMessageBox.warning(self, "输入图像不存在", input_path)
            self._log(f"[ERROR] 输入图像不存在: {input_path}")
            return
        output_dir = self._ensure_output_dir()
        if not output_dir or not self._validate_algorithm_device():
            return
        options = self._current_options()
        options["output_dir"] = output_dir
        self.radio_single.setChecked(True)
        self._prepare_run("single")
        self.gen_worker = GenerationWorker(input_path, self.lbl_gt_path.text().strip(), options)
        self._connect_worker(self.gen_worker)
        self.gen_worker.start()

    def _on_generate_batch(self):
        if self._is_generation_running():
            return
        input_dir = self.lbl_input_folder.text().strip()
        if not input_dir:
            QMessageBox.warning(self, "缺少输入文件夹", "请先选择批量输入文件夹。")
            self._log("[ERROR] 没有选择输入文件夹")
            return
        if not os.path.isdir(input_dir):
            QMessageBox.warning(self, "输入文件夹不存在", input_dir)
            self._log(f"[ERROR] 输入文件夹不存在: {input_dir}")
            return
        try:
            input_files = scan_image_files(input_dir)
        except Exception as exc:
            QMessageBox.warning(self, "输入文件夹无效", str(exc))
            self._log(f"[ERROR] 输入文件夹扫描失败: {exc}")
            return
        if not input_files:
            QMessageBox.warning(self, "没有可处理图像", "输入文件夹中没有支持格式的图像。")
            self._log("[ERROR] 输入文件夹中没有支持格式的图像")
            return
        label_dir = self.lbl_gt_folder.text().strip()
        if label_dir and not os.path.isdir(label_dir):
            QMessageBox.warning(self, "标签文件夹不存在", label_dir)
            self._log(f"[ERROR] 标签文件夹不存在: {label_dir}")
            return
        output_dir = self._ensure_output_dir()
        if not output_dir or not self._validate_algorithm_device():
            return
        options = self._current_options()
        options["output_dir"] = output_dir
        self.radio_batch.setChecked(True)
        self._prepare_run("batch")
        self.gen_worker = BatchGenerationWorker(input_files, input_dir, label_dir, options)
        self._connect_worker(self.gen_worker)
        self.gen_worker.start()

    def _connect_worker(self, worker):
        worker.log_message.connect(self._log)
        worker.progress_changed.connect(self._on_progress_changed)
        worker.preview_ready.connect(self._on_preview_ready)
        worker.result_ready.connect(self._on_generation_result)
        worker.error_occurred.connect(self._on_generation_error)
        worker.finished.connect(self._on_worker_finished)
        worker.finished.connect(worker.deleteLater)

    def _prepare_run(self, mode):
        self._reset_progress()
        self.generated_image = None
        self.compare_image = None
        self.lbl_generation_status.setText("生成中" if mode == "single" else "批量生成中")
        self.lbl_current_file.setText("-")
        self.lbl_output_image.setText("生成中...")
        self.lbl_compare_image.setText("等待生成对比图")
        self.metric_labels["save_path"].setText("等待保存")
        self.metric_labels["report_path"].setText("等待导出")
        self._set_controls_enabled(False)

    def _set_controls_enabled(self, enabled):
        for widget in (
            self.radio_single,
            self.radio_batch,
            self.btn_select_input,
            self.btn_select_input_folder,
            self.btn_select_gt,
            self.btn_select_gt_folder,
            self.btn_select_output_dir,
            self.combo_algorithm,
            self.combo_device,
            self.edit_filename_prefix,
            self.chk_compute_metrics,
            self.chk_save_compare,
            self.chk_export_report,
            self.chk_export_docx,
            self.btn_start_single,
            self.btn_start_batch,
        ):
            widget.setEnabled(enabled)
        self.btn_cancel.setEnabled(not enabled)
        if enabled:
            self._on_mode_changed()

    def _on_cancel_task(self):
        if self.gen_worker is None:
            return
        self.gen_worker.request_cancel()
        self.btn_cancel.setEnabled(False)
        self.lbl_generation_status.setText("取消中")
        self._log("[CANCEL] 已请求取消任务，当前图像处理完成后停止")

    def _on_progress_changed(self, stats):
        self.progress_bar.setValue(int(stats.get("percent", 0)))
        completed = stats.get("completed", 0)
        total = stats.get("total", 0)
        self.progress_labels["completed"].setText(f"{completed}/{total}")
        self.progress_labels["success"].setText(str(stats.get("success", 0)))
        self.progress_labels["failed"].setText(str(stats.get("failed", 0)))
        self.progress_labels["current_elapsed"].setText(f"{stats.get('current_elapsed', 0):.3f} s")
        self.progress_labels["total_elapsed"].setText(f"{stats.get('total_elapsed', 0):.3f} s")
        self.progress_labels["avg_elapsed"].setText(f"{stats.get('avg_elapsed', 0):.3f} s")
        self.lbl_current_file.setText(stats.get("current_file") or "-")
        self.lbl_generation_status.setText(stats.get("status") or "运行中")

    def _on_preview_ready(self, input_image, generated_image, label_image, compare_image):
        if input_image is not None:
            self.input_image = input_image
            show_image_on_label(self.lbl_input_image, input_image)
        if generated_image is not None:
            self.generated_image = generated_image
            show_image_on_label(self.lbl_output_image, generated_image)
        if label_image is not None:
            self.gt_image = label_image
            show_image_on_label(self.lbl_gt_image, label_image)
        else:
            self.gt_image = None
            self.lbl_gt_image.clear()
            if self.lbl_gt_folder.text().strip():
                self.lbl_gt_image.setText("当前图像未匹配标签")
            elif self.lbl_gt_path.text().strip():
                self.lbl_gt_image.setText("标签未加载")
            else:
                self.lbl_gt_image.setText("未选择标签")
        if compare_image is not None:
            self.compare_image = compare_image
            show_image_on_label(self.lbl_compare_image, compare_image)

    def _on_generation_result(self, result):
        self._set_controls_enabled(True)
        self.lbl_generation_status.setText("已取消" if result.get("canceled") else "已完成")
        record = result.get("record")
        if record:
            self._display_record_metrics(record)
            self.metric_labels["save_path"].setText(record.get("output_path") or result.get("run_dir", ""))
            self.metric_labels["report_path"].setText(result.get("report_txt") or result.get("params_path") or "未导出")
        else:
            summary = result.get("summary", {})
            self.metric_labels["save_path"].setText(result.get("run_dir", ""))
            self.metric_labels["report_path"].setText(result.get("report_txt", ""))
            for batch_record in reversed(result.get("records", [])):
                if batch_record.get("status") == "success":
                    self._display_record_metrics(batch_record)
                    break
            if summary:
                self.progress_labels["completed"].setText(f"{summary.get('completed', 0)}/{summary.get('total', 0)}")
        self._log("[INFO] 任务结果已回写到界面")

    def _display_record_metrics(self, record):
        for key in ("psnr", "ssim", "l1", "lipips"):
            value = record.get(key)
            if value not in ("", None):
                self.metric_labels[key].setText(value)
            elif key == "lipips" and any(record.get(metric_key) for metric_key in ("psnr", "ssim", "l1")):
                self.metric_labels[key].setText("无法计算")
            else:
                self.metric_labels[key].setText("未提供标签")

    def _on_generation_error(self, error_msg):
        self._set_controls_enabled(True)
        self.lbl_generation_status.setText("生成失败")
        self.lbl_output_image.setText("生成失败，请查看运行日志")
        self.metric_labels["save_path"].setText("未保存")
        self._log(f"[ERROR] 生成失败: {error_msg}")
        QMessageBox.critical(self, "生成失败", f"生成过程中发生错误:\n\n{error_msg}")

    def _on_worker_finished(self):
        worker = self.sender()
        if self.gen_worker is worker:
            self.gen_worker = None
        self._set_controls_enabled(True)

    def _is_generation_running(self):
        return self.gen_worker is not None and self.gen_worker.isRunning()

    def _warn_if_running(self):
        if not self._is_generation_running():
            return False
        QMessageBox.information(self, "任务运行中", "当前任务尚未结束，请等待完成或取消后再修改输入。")
        return True

    def _reset_progress(self):
        if hasattr(self, "progress_bar"):
            self.progress_bar.setValue(0)
        for key, label in self.progress_labels.items():
            if key == "completed":
                label.setText("0/0")
            elif key in ("success", "failed"):
                label.setText("0")
            else:
                label.setText("0.000 s")
        for key in ("psnr", "ssim", "l1", "lipips"):
            if key in self.metric_labels:
                self.metric_labels[key].setText("未提供标签")

    def _on_open_save_dir(self):
        save_dir = self.lbl_output_dir.text().strip() or DEFAULT_OUTPUT_DIR
        os.makedirs(save_dir, exist_ok=True)
        if sys.platform == "win32":
            os.startfile(save_dir)
        else:
            subprocess.Popen(["xdg-open", save_dir])

    def _log(self, message):
        self.log_text.append(str(message))
        if self.log_signal:
            self.log_signal.emit(str(message))

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self.input_image is not None:
            show_image_on_label(self.lbl_input_image, self.input_image)
        if self.generated_image is not None:
            show_image_on_label(self.lbl_output_image, self.generated_image)
        if self.gt_image is not None:
            show_image_on_label(self.lbl_gt_image, self.gt_image)
        if self.compare_image is not None:
            show_image_on_label(self.lbl_compare_image, self.compare_image)
